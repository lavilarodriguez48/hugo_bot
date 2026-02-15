import os
import json
from docx import Document

INPUT_DIR = "data/raw"
OUTPUT_FILE = "data/paragraphs.jsonl"

def extract_from_paragraph(paragraph):
    """Extrae texto y detecta formato (negrita, color, resaltado)."""
    text = paragraph.text.strip()
    if not text:
        return None

    bold = False
    underline = False
    highlighted = False
    colored = False

    for run in paragraph.runs:
        if run.bold:
            bold = True
        if run.underline:
            underline = True
        if run.font.highlight_color:
            highlighted = True
        if run.font.color and run.font.color.rgb:
            colored = True

    return {
        "text": text,
        "bold": bold,
        "underline": underline,
        "highlighted": highlighted,
        "colored": colored
    }

def extract_from_table(table):
    """Extrae texto y formato de cada celda de una tabla."""
    rows_data = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                data = extract_from_paragraph(paragraph)
                if data:
                    rows_data.append(data)
    return rows_data

def main():
    with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
        for filename in os.listdir(INPUT_DIR):
            if not filename.endswith(".docx"):
                continue

            full_path = os.path.join(INPUT_DIR, filename)
            doc = Document(full_path)

            # Extraer párrafos normales
            for p in doc.paragraphs:
                data = extract_from_paragraph(p)
                if data:
                    out.write(json.dumps(data) + "\n")

            # Extraer tablas
            for table in doc.tables:
                rows = extract_from_table(table)
                for row in rows:
                    out.write(json.dumps(row) + "\n")

if __name__ == "__main__":
    main()



