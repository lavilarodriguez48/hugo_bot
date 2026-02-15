from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os

OUTPUT_DIR = "data/raw"

def save_doc(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print("Creado:", path)

# -------------------------------
# 1. Examen con negrita y colores
# -------------------------------
doc1 = Document()
p = doc1.add_paragraph()
p.add_run("1. ¿Qué es la fotosíntesis?").bold = True
doc1.add_paragraph("A) Proceso por el cual las plantas respiran")
r = doc1.add_paragraph().add_run("B) Proceso por el cual las plantas producen alimento")
r.bold = True
r.font.color.rgb = RGBColor(0, 128, 0)
doc1.add_paragraph("C) Movimiento de nutrientes en el suelo")
doc1.add_paragraph("D) Ninguna de las anteriores")
save_doc(doc1, "examen_negrita_colores.docx")

# -----------------------------------------
# 2. Examen con ✔, resaltado y varias correctas
# -----------------------------------------
doc2 = Document()
doc2.add_paragraph("1- Selecciona todas las opciones correctas:")
r = doc2.add_paragraph().add_run("✔ A) Tienen glándulas mamarias")
r.bold = True
r.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc2.add_paragraph("✔ B) Son de sangre caliente")
doc2.add_paragraph("C) Ponen huevos")
doc2.add_paragraph("D) No respiran oxígeno")
save_doc(doc2, "examen_resaltado_varias_correctas.docx")

# -------------------------------
# 3. Examen con tablas
# -------------------------------
doc3 = Document()
table = doc3.add_table(rows=3, cols=2)
table.style = "Table Grid"

table.cell(0,0).text = "1_ ¿Cuál es la capital de Francia?"
table.cell(0,1).text = "A) Roma\nB) París\nC) Berlín\nD) Lisboa"

table.cell(1,0).text = "2) Indica la correcta:"
table.cell(1,1).text = "A) El agua hierve a 50°C\nB) El agua no hierve\nC) El agua hierve a 100°C\nD) El agua hierve a 200°C"

save_doc(doc3, "examen_tablas.docx")

# -------------------------------
# 4. Verdadero/Falso
# -------------------------------
doc4 = Document()
doc4.add_paragraph("1 ) Verdadero o falso: La Tierra gira alrededor del Sol.")
doc4.add_paragraph("A . Verdadero")
r = doc4.add_paragraph().add_run("B . Falso")
r.bold = True
save_doc(doc4, "examen_verdadero_falso.docx")

# -------------------------------
# 5. Examen con opciones raras
# -------------------------------
doc5 = Document()
doc5.add_paragraph("1: ¿Qué órgano bombea la sangre?")
doc5.add_paragraph("a) Pulmones")
doc5.add_paragraph("a- Corazón")
doc5.add_paragraph("a_ Estómago")
doc5.add_paragraph("A . Hígado")
save_doc(doc5, "examen_opciones_raras.docx")

# -------------------------------
# 6. Examen con mezclas extremas
# -------------------------------
doc6 = Document()
p = doc6.add_paragraph()
p.add_run("1 . Señala la correcta:").font.color.rgb = RGBColor(255,0,0)
doc6.add_paragraph("A ) Opción incorrecta")
r = doc6.add_paragraph().add_run("A- Opción correcta")
r.bold = True
r.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc6.add_paragraph("A_ Otra incorrecta")
save_doc(doc6, "examen_mezcla_extrema.docx")

# -------------------------------
# 7. Examen con varias correctas sin avisar
# -------------------------------
doc7 = Document()
doc7.add_paragraph("1_ ¿Qué animales son mamíferos?")
r = doc7.add_paragraph().add_run("A) Perro")
r.bold = True
r = doc7.add_paragraph().add_run("B) Gato")
r.font.color.rgb = RGBColor(0,0,255)
doc7.add_paragraph("C) Tortuga")
doc7.add_paragraph("D) Serpiente")
save_doc(doc7, "examen_varias_correctas_sin_avisar.docx")

# -------------------------------
# 8. Examen con formatos mixtos
# -------------------------------
doc8 = Document()
doc8.add_paragraph("1- Completa la frase:")
r = doc8.add_paragraph().add_run("A) El agua es un ")
r.bold = True
doc8.add_paragraph("A) líquido")
doc8.add_paragraph("A . sólido")
doc8.add_paragraph("A_ gas")
save_doc(doc8, "examen_formatos_mixtos.docx")
