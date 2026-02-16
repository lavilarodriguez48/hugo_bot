from docx import Document
import re

def es_pregunta(texto):
    # 1. / 1) / 1- / 1º ...
    return re.match(r"\s*\d+[\.\)\-\º]\s+", texto) is not None

def es_opcion(texto):
    # a) / b. / c) / d- / A) / B.
    return re.match(r"\s*[a-dA-D][\)\.\-\:]\s+", texto) is not None

def limpiar_prefijo_pregunta(texto):
    return re.sub(r"^\s*\d+[\.\)\-\º]\s+", "", texto).strip()

def limpiar_prefijo_opcion(texto):
    return re.sub(r"^\s*[a-dA-D][\)\.\-\:]\s+", "", texto).strip()

def parrafo_es_correcto(paragraph):
    # Cualquier formato “especial” lo consideramos correcto
    for run in paragraph.runs:
        if run.bold:
            return True
        if run.underline:
            return True
        if run.font.highlight_color is not None:
            return True
        if run.font.color and run.font.color.rgb is not None:
            return True
        # Marcas tipo X), *, ✔ dentro del texto
        if re.search(r"\b[Xx\*✔]\b", run.text):
            return True
    return False

def procesar_parrafos(paragraphs, preguntas):
    actual = None

    for p in paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        if es_pregunta(texto):
            # cerramos la anterior
            if actual and actual["opciones"]:
                # si no se detectó correcta, no la dejamos colgada
                if actual["correcta"] is None and len(actual["opciones"]) == 1:
                    actual["correcta"] = 0
                preguntas.append(actual)

            actual = {
                "enunciado": limpiar_prefijo_pregunta(texto),
                "opciones": [],
                "correcta": None
            }

        elif es_opcion(texto) and actual is not None:
            opcion_limpia = limpiar_prefijo_opcion(texto)
            idx = len(actual["opciones"])
            actual["opciones"].append(opcion_limpia)

            if parrafo_es_correcto(p):
                actual["correcta"] = idx

    if actual and actual["opciones"]:
        if actual["correcta"] is None and len(actual["opciones"]) == 1:
            actual["correcta"] = 0
        preguntas.append(actual)

def extraer_preguntas(docx_file):
    doc = Document(docx_file)
    preguntas = []

    # 1) Procesar párrafos “sueltos”
    procesar_parrafos(doc.paragraphs, preguntas)

    # 2) Procesar tablas (cada celda como si fuera un párrafo)
    for table in doc.tables:
        celdas = []
        for row in table.rows:
            for cell in row.cells:
                # Evitar duplicados de celdas compartidas
                if cell not in celdas:
                    celdas.append(cell)

        paragraphs = []
        for cell in celdas:
            paragraphs.extend(cell.paragraphs)

        procesar_parrafos(paragraphs, preguntas)

    # Filtrar preguntas sin correcta o sin opciones
    preguntas_limpias = []
    for p in preguntas:
        if not p["opciones"]:
            continue
        if p["correcta"] is None:
            # si no hay correcta, por seguridad no la incluimos
            continue
        if p["correcta"] < 0 or p["correcta"] >= len(p["opciones"]):
            continue
        preguntas_limpias.append(p)

    return preguntas_limpias
